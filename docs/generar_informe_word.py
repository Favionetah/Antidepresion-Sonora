import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = "docs"

def shade_cell(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, bold=False, color=None, size=9):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = RGBColor(*color)

def crear_informe():
    doc = Document()

    # ── Estilos globales ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ═══════════════ PORTADA ═══════════════
    for _ in range(6):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PROYECTO FINAL")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e)
    run.bold = True

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ELABORACIÓN DE UNA IA FUNCIONAL")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x39, 0x49, 0xab)

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Chatbot SBC Musicoterapia")
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Sistema Basado en Conocimiento para Diagnóstico\n"
                     "de Estrés y Prescripción de Musicoterapia")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x54, 0x6e, 0x7a)

    for _ in range(4):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Arquitectura: FSM (Finite State Machine) con 29 Nodos\n"
                     "Motor de Reglas: Base de Conocimiento Singleton + Deteccion de Emociones + Matcher de Intenciones\n"
                     "Integracion: Spotify API via OAuth 2.0\n"
                     "Interfaz: CustomTkinter (Desktop) + python-telegram-bot\n"
                     "Version: 2.0 (Fase 2 - Integracion Hibrida)")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x54, 0x6e, 0x7a)

    doc.add_page_break()

    # ═══════════════ ÍNDICE ═══════════════
    doc.add_heading("Índice", level=1)
    items = [
        "1. Introducción",
        "2. Objetivos",
        "3. Arquitectura del Sistema",
        "4. Árbol de Decisión (29 Nodos)",
        "5. Tabla Completa de Diagnósticos",
        "6. Diagrama de Flujo del Motor FSM",
        "7. Capturas de Pantalla",
        "8. Conclusiones",
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ═══════════════ 1. INTRODUCCIÓN ═══════════════
    doc.add_heading("1. Introducción", level=1)
    doc.add_paragraph(
        "El presente documento describe el desarrollo de un Sistema Basado en Conocimiento (SBC) "
        "implementado como un chatbot interactivo para el diagnóstico de estrés y la prescripción "
        "personalizada de musicoterapia a través de Spotify. El sistema utiliza un motor de Máquina "
        "de Estados Finitos (FSM) con 29 nodos que guían al usuario a través de una entrevista "
        "estructurada para identificar su estado emocional y recomendar la terapia musical más adecuada."
    )
    doc.add_paragraph(
        "El proyecto integra múltiples componentes de inteligencia artificial: detección de emociones "
        "mediante análisis de palabras clave en español, coincidencia de intenciones por similitud "
        "semántica con expansión de sinónimos y stemming, y un sistema de redirección emocional que "
        "permite adaptar el flujo de la conversación según el estado anímico detectado del usuario."
    )
    doc.add_paragraph(
        "La aplicación se despliega en dos interfaces: una aplicación de escritorio con CustomTkinter "
        "y un bot de Telegram, ambas compartiendo el mismo núcleo de conocimiento y lógica FSM."
    )

    # ═══════════════ 2. OBJETIVOS ═══════════════
    doc.add_heading("2. Objetivos", level=1)
    doc.add_heading("Objetivo General", level=2)
    doc.add_paragraph(
        "Desarrollar un chatbot funcional basado en reglas y lógica de estados que sea capaz de "
        "diagnosticar problemas de estrés y recomendar musicoterapia personalizada mediante la "
        "integración con la API de Spotify."
    )
    doc.add_heading("Objetivos Específicos", level=2)
    objetivos = [
        "Implementar un árbol de decisión con al menos 25 nodos que cubra distintos tipos de estrés.",
        "Diseñar un motor FSM que gestione las transiciones entre estados y el historial de sesión.",
        "Integrar un módulo de detección de emociones para adaptar dinámicamente el flujo.",
        "Implementar un matcher de intenciones para procesar texto libre del usuario.",
        "Conectar con la API de Spotify para buscar y reproducir playlists terapéuticas.",
        "Construir una interfaz gráfica de escritorio amigable con CustomTkinter.",
        "Documentar todo el árbol de decisión con diagramas visuales.",
    ]
    for obj in objetivos:
        p = doc.add_paragraph(obj, style='List Bullet')

    # ═══════════════ 3. ARQUITECTURA ═══════════════
    doc.add_heading("3. Arquitectura del Sistema", level=1)
    doc.add_heading("3.1 Componentes del Sistema", level=2)
    componentes = [
        ("base_conocimiento.py", "Singleton que contiene los 29 nodos con mensajes, opciones y transiciones."),
        ("motor_fsm.py", "Máquina de estados finitos que gestiona transiciones y redirección emocional."),
        ("contexto.py", "Dataclass SesionSBC con estado de sesión, historial y recomendaciones."),
        ("emociones.py", "Detector de emociones en español con 7 categorías y modificadores de intensidad."),
        ("intenciones.py", "Matcher de intenciones con expansión de sinónimos, stemming y n-gramas."),
        ("frases.py", "Banco de frases templatizadas con selección aleatoria y deduplicación."),
        ("spotify.py", "Integración con API de Spotify para búsqueda y reproducción de playlists."),
        ("desktop_gui.py", "Interfaz gráfica con CustomTkinter (ventana 520x720)."),
        ("telegram_bot.py", "Interfaz para Telegram con manejo de comandos y conversaciones."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Componente"
    hdr[1].text = "Descripción"
    for i, (nombre, desc) in enumerate(componentes):
        row = table.add_row().cells
        row[0].text = nombre
        row[1].text = desc

    doc.add_heading("3.2 Flujo de una Conversación", level=2)
    pasos = [
        ("1. Inicio", "El usuario inicia la aplicación y recibe un mensaje de bienvenida aleatorio."),
        ("2. Pregunta", "El bot presenta una pregunta con 2-3 opciones en forma de botones."),
        ("3. Entrada", "El usuario puede hacer clic en un botón o escribir texto libre."),
        ("4. Procesamiento", "Si es texto, se detecta emoción y se empareja la intención con las opciones disponibles."),
        ("5. Redirección", "Si la emoción no coincide con la esperada, se sugiere una rama alternativa."),
        ("6. Transición", "El motor FSM avanza al siguiente nodo según la opción seleccionada."),
        ("7. Diagnóstico", "Al alcanzar un nodo hoja, se muestra el diagnóstico y se busca la playlist."),
        ("8. Reproducción", "El usuario puede abrir la playlist en Spotify o reproducirla directamente."),
    ]
    for titulo, desc in pasos:
        p = doc.add_paragraph()
        run = p.add_run(f"{titulo}: ")
        run.bold = True
        p.add_run(desc)

    # ═══════════════ 4. ÁRBOL DE DECISIÓN ═══════════════
    doc.add_heading("4. Árbol de Decisión (29 Nodos)", level=1)
    doc.add_paragraph(
        "El árbol de decisión consta de 29 nodos distribuidos en 5 niveles. "
        "Cada nodo interno presenta una pregunta con opciones que derivan a nodos hijos. "
        "Los nodos hoja (13 en total) contienen el diagnóstico final y la recomendación de playlist."
    )

    ruta_img = os.path.join(OUT_DIR, "arbol_decision.png")
    if os.path.exists(ruta_img):
        doc.add_picture(ruta_img, width=Inches(6.2))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Leyenda de colores: Azul oscuro (Raíz) → Azul medio (Tipo de síntoma) → "
        "Azul claro (Clasificación) → Lavanda (Enfoque terapéutico) → Verde (Diagnóstico final)."
    )

    doc.add_heading("4.1 Niveles del Árbol", level=2)
    niveles_tabla = [
        ("Nivel 0", "Raíz", "1", "N01"),
        ("Nivel 1", "Tipo de Síntoma", "3", "N02-N04"),
        ("Nivel 2", "Clasificación del Estrés", "6", "N05-N10"),
        ("Nivel 3", "Enfoque Terapéutico", "6", "N11-N16"),
        ("Nivel 4", "Intervención / Hoja", "13", "N17-N25, N28-N31"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(["Nivel", "Descripción", "Cant. Nodos", "IDs"]):
        hdr[i].text = h
    for row_data in niveles_tabla:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    doc.add_heading("4.2 Tabla Completa de Nodos", level=2)
    nodos_tabla = [
        ("N01", "Raíz", "Síntoma predominante", "3 opciones", "No"),
        ("N02", "Físico", "Agudo vs Crónico", "2 opciones", "No"),
        ("N03", "Mental", "Bloqueo vs Hiperactividad", "2 opciones", "No"),
        ("N04", "Emocional", "Frustración vs Tristeza", "2 opciones", "No"),
        ("N05", "Estrés Agudo", "Enfoque de manejo", "2 opciones", "No"),
        ("N06", "Burnout", "Sueño vs Relajación", "2 opciones", "No"),
        ("N07", "Bloqueo Mental", "Abordaje", "2 opciones", "No"),
        ("N08", "Overthinking", "Enfoque calmante", "2 opciones", "No"),
        ("N09", "Frustración", "Canalización", "2 opciones", "No"),
        ("N10", "Melancolía", "Ambiente musical", "2 opciones", "No"),
        ("N11", "Frecuencias", "Solfeggio: 528Hz vs 432Hz", "2 opciones", "No"),
        ("N12", "Naturaleza", "Lluvia vs Ruido marrón", "2 opciones", "No"),
        ("N13", "Ambiente denso", "Dark Ambient vs Deep Sleep", "2 opciones", "No"),
        ("N14", "Binaurales", "Activar ondas alpha", "1 opción", "No"),
        ("N15", "Lo-Fi Chillhop", "Sumergirse en ritmos", "1 opción", "No"),
        ("N16", "Post-Rock", "Viaje catártico", "1 opción", "No"),
        ("N17", "528Hz", "Estrés Agudo Alta Activación", "Hoja", "Sí"),
        ("N18", "432Hz", "Estrés Agudo Calma Profunda", "Hoja", "Sí"),
        ("N19", "Lluvia/Tormenta", "Crisis de Ansiedad", "Hoja", "Sí"),
        ("N20", "Ruido Marrón", "Ansiedad Aguda", "Hoja", "Sí"),
        ("N21", "Dark Ambient", "Agotamiento Crónico", "Hoja", "Sí"),
        ("N22", "Deep Sleep", "Insomnio Crónico", "Hoja", "Sí"),
        ("N23", "Ondas Alpha", "Bloqueo Cognitivo", "Hoja", "Sí"),
        ("N24", "Lo-Fi Beats", "Ansiedad Cognitiva", "Hoja", "Sí"),
        ("N25", "Post-Rock", "Frustración Catarsis", "Hoja", "Sí"),
        ("N28", "Lo-Fi Creativo", "Bloqueo Flujo Creativo", "Hoja", "Sí"),
        ("N29", "432Hz Calmante", "Hiperactividad Mental", "Hoja", "Sí"),
        ("N30", "Solfeggio Energía", "Fatiga Crónica", "Hoja", "Sí"),
        ("N31", "Épica Instrumental", "Frustración Liberación", "Hoja", "Sí"),
    ]
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(["Nodo", "Categoría", "Descripción", "Tipo", "Hoja"]):
        hdr[i].text = h
    for row_data in nodos_tabla:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            set_cell_text(row[i], val, size=8)
            if row_data[4] == "Sí":
                shade_cell(row[i], "E8F5E9")

    # ═══════════════ 5. DIAGNÓSTICOS ═══════════════
    doc.add_heading("5. Tabla Completa de Diagnósticos", level=1)
    doc.add_paragraph(
        "Cada nodo hoja contiene un diagnóstico único con su respectiva recomendación "
        "terapéutica y una playlist de Spotify asociada."
    )

    diagnosticos = [
        ("N17", "Estrés Agudo con Alta Activación Fisiológica",
         "528Hz Solfeggio Healing Frequency", "528Hz Reparación y Sanación",
         "Reducción de cortisol, reparación celular"),
        ("N18", "Estrés Agudo con Necesidad de Calma Profunda",
         "432Hz Calming Deep Relaxation", "432Hz Calma Profunda",
         "Armonización, calma profunda"),
        ("N19", "Crisis de Ansiedad con Necesidad de Escudo Acústico",
         "Rain Thunderstorm Sounds", "Lluvia y Tormenta para Dormir",
         "Enmascaramiento de ruido ambiental"),
        ("N20", "Ansiedad Aguda con Ruido Mental Persistente",
         "Brown Noise Deep Focus Sleep", "Ruido Marrón Profundo",
         "Aislamiento sensorial, ruido marrón"),
        ("N21", "Agotamiento Crónico con Desconexión Profunda",
         "Dark Ambient Drone Space Music", "Dark Ambient Espacial",
         "Despersonalización del estrés"),
        ("N22", "Insomnio por Estrés Crónico con Fatiga Profunda",
         "Deep Sleep Space Music Delta Waves", "Deep Sleep Space Music",
         "Inducción de ondas delta (sueño)"),
        ("N23", "Bloqueo Mental por Sobrecarga Cognitiva",
         "Binaural Beats Alpha Waves Focus", "Ondas Alpha para Enfoque",
         "Sincronización de ondas alpha 8-12Hz"),
        ("N24", "Ansiedad Cognitiva con Hiperactividad Mental",
         "Lo-Fi Chillhop Beats Study Relax", "Lo-Fi Beats sin Letra",
         "BPM 60-80 para calma activa"),
        ("N25", "Frustración Acumulada con Catarsis Emocional",
         "Post-Rock Instrumental Cinematic", "Post-Rock Instrumental Catártico",
         "Crescendos emocionales liberadores"),
        ("N28", "Bloqueo Mental — Activación de Flujo Creativo",
         "Lo-Fi Hip Hop Focus Beats", "Lo-Fi Foco Creativo",
         "Ritmo mental para creatividad"),
        ("N29", "Hiperactividad Mental — Calma y Silencio Interior",
         "432Hz Calming Music Peaceful", "Frecuencias Calmantes 432Hz",
         "Aquietamiento del sistema nervioso"),
        ("N30", "Fatiga Crónica — Reparación Energética",
         "Solfeggio Frequencies 528Hz 432Hz", "Frecuencias de Reparación Energética",
         "Restauración del equilibrio energético"),
        ("N31", "Frustración — Liberación de Tensión Acumulada",
         "Epic Powerful Instrumental Orchestra", "Épica Instrumental Liberadora",
         "Orquestaciones grandiosas, empoderamiento"),
    ]
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(["Nodo", "Diagnóstico", "Query Spotify", "Playlist", "Efecto Terapéutico"]):
        hdr[i].text = h
    for row_data in diagnosticos:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            set_cell_text(row[i], val, size=7.5)

    # ═══════════════ 6. DIAGRAMA FSM ═══════════════
    doc.add_heading("6. Diagrama de Flujo del Motor FSM", level=1)
    doc.add_paragraph(
        "El siguiente diagrama muestra el flujo completo del motor de Máquina de Estados Finitos (FSM) "
        "que gobierna el comportamiento del chatbot. Incluye los pasos de detección de emociones, "
        "coincidencia de intenciones, validación de redirección emocional y generación del diagnóstico."
    )
    ruta_fsm = os.path.join(OUT_DIR, "diagrama_flujo_fsm.png")
    if os.path.exists(ruta_fsm):
        doc.add_picture(ruta_fsm, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "El motor FSM implementa un patrón Singleton para el contexto de sesión, "
        "permitiendo que múltiples interacciones (desktop y Telegram) compartan la misma lógica."
    )

    # ═══════════════ 7. CAPTURAS ═══════════════
    doc.add_heading("7. Capturas de Pantalla", level=1)
    doc.add_heading("7.1 Pantalla de Inicio", level=2)
    doc.add_paragraph(
        "La ventana principal muestra el mensaje de bienvenida del chatbot, la primera pregunta "
        "con tres opciones en forma de botones, y el campo de entrada de texto libre en la parte inferior."
    )
    ruta_inicio = os.path.join(OUT_DIR, "captura_pantalla_inicio.png")
    if os.path.exists(ruta_inicio):
        doc.add_picture(ruta_inicio, width=Inches(3.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("7.2 Pantalla de Diagnóstico", level=2)
    doc.add_paragraph(
        "Durante la conversación, el usuario responde a las preguntas del bot mediante opciones "
        "o texto libre. Al llegar a un nodo hoja, se muestra el diagnóstico con la recomendación "
        "terapéutica y el nombre de la playlist recomendada."
    )
    ruta_diag = os.path.join(OUT_DIR, "captura_pantalla_diagnostico.png")
    if os.path.exists(ruta_diag):
        doc.add_picture(ruta_diag, width=Inches(3.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("7.3 Pantalla de Resultado Spotify", level=2)
    doc.add_paragraph(
        "Una vez completado el diagnóstico, la aplicación se conecta a la API de Spotify para buscar "
        "la playlist recomendada. El usuario puede abrir la playlist en el navegador o reproducirla "
        "directamente si está autenticado con una cuenta Premium."
    )
    ruta_spot = os.path.join(OUT_DIR, "captura_pantalla_spotify.png")
    if os.path.exists(ruta_spot):
        doc.add_picture(ruta_spot, width=Inches(3.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ═══════════════ 8. CONCLUSIONES ═══════════════
    doc.add_heading("8. Conclusiones", level=1)
    conclusiones = [
        "Se implementó exitosamente un chatbot funcional con 29 nodos que cubre 13 diagnósticos "
        "diferenciados de estrés, cumpliendo con el requisito mínimo de 25 nodos.",
        "La arquitectura FSM con base de conocimiento Singleton demostró ser robusta y escalable, "
        "permitiendo la integración simultánea con interfaces de escritorio y Telegram.",
        "El módulo de detección de emociones y redirección dinámica agrega un valor significativo "
        "a la experiencia del usuario, permitiendo que el chatbot se adapte al estado anímico real "
        "de la persona más allá de las opciones predefinidas.",
        "La integración con la API de Spotify a través de OAuth 2.0 permite ofrecer recomendaciones "
        "musicales personalizadas y funcionales, no solo teóricas.",
        "El matcher de intenciones con procesamiento de lenguaje natural en español (stemming, "
        "expansión de sinónimos, n-gramas) permite una interacción más natural, aceptando texto "
        "libre además de clics en botones.",
        "El sistema de frases templatizadas con selección aleatoria y deduplicación evita la "
        "repetición excesiva y mejora la percepción de naturalidad del chatbot.",
        "Se recomienda como trabajo futuro: ampliar el árbol a dominios adicionales, mejorar "
        "el modelo de detección emocional con machine learning, y agregar análisis de sentimiento "
        "en tiempo real sobre la retroalimentación del usuario.",
    ]
    for c in conclusiones:
        p = doc.add_paragraph(c, style='List Number')

    # ── Guardar ──
    ruta_docx = os.path.join(OUT_DIR, "Informe_Proyecto_Chatbot_SBC.docx")
    doc.save(ruta_docx)
    print(f"OK: {ruta_docx} generado")

if __name__ == "__main__":
    os.makedirs(OUT_DIR, exist_ok=True)
    crear_informe()
